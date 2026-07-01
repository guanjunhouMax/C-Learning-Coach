# -*- coding: utf-8 -*-
"""生成C语言答疑助手项目提案报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_report():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # ===== 封面 =====
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《人工智能基础与应用》课程考核项目提案")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3a)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("项目名称：基于大语言模型的C语言编程学习助手")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x3a, 0x5a, 0x8a)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("项目成员（个人项目）")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ===== 摘要 =====
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "本项目旨在设计并实现一个基于大语言模型的C语言编程学习助手——「C语言学习小教练」。"
        "该助手以Web应用的形式呈现，能够回答用户在C语言编程学习中遇到的各种问题，包括基础语法、"
        "控制流、数组、指针、函数、字符串、结构体、文件操作、内存管理、预处理指令等多个核心知识点。"
        "系统支持两种运行模式：一是通过大语言模型API（如OpenAI GPT系列）提供高质量、智能化的对话式答疑；"
        "二是在无API密钥的情况下，内置了涵盖C语言核心知识点的本地知识库作为备选方案。"
        "本项目结合了计算机网络专业背景，涉及API调用、HTTP通信协议、前端-后端数据交互等技术，"
        "是AI技术与编程教育相结合的一次有益尝试。"
    )

    # ===== 背景与意义 =====
    doc.add_heading("背景与意义", level=1)
    doc.add_paragraph(
        "C语言作为计算机科学领域的经典编程语言，是许多高校计算机专业的第一门编程语言课程。"
        "然而，对于编程初学者来说，学习C语言常常会遇到诸多困难：指针概念抽象难懂、内存管理容易出错、"
        "编译错误信息晦涩难懂、遇到问题得不到及时解答等。传统的学习方式依赖于教材或论坛提问，效率低下且响应不及时。"
    )
    doc.add_paragraph(
        "本项目正是为了解决这一问题而设计。通过构建一个AI驱动的C语言编程学习助手，为学习者提供"
        "7×24小时的即时答疑服务。项目的特色与价值体现在以下几个方面："
    )
    items = [
        "即时性：利用大语言模型API实现秒级响应，学习者无需等待即可获得解答。",
        "针对性：专注C语言编程领域，涵盖指针、内存管理等难点知识，回答包含大量可运行的代码示例。",
        "双模式保障：即使没有API密钥，内置知识库也能覆盖C语言核心知识点，确保基本功能可用。",
        "专业性结合：项目涉及网络API调用、前后端交互、数据序列化等技术，与计算机网络专业课程紧密相关。",
        "可扩展性：基于Flask框架构建，便于后续添加练习系统、代码运行等更多功能。"
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # ===== 核心功能 =====
    doc.add_heading("核心功能", level=1)
    features = [
        ("智能对话答疑", "基于大语言模型API的智能对话功能，能够理解用户自然语言问题，给出准确、详细的解答，并附带C语言代码示例。"),
        ("本地知识库问答", "内置涵盖C语言基础语法、控制流、数组、指针、函数、字符串、结构体、文件操作、内存管理、预处理指令等11个核心主题的知识库，支持关键词匹配的离线问答模式。"),
        ("代码格式化展示", "回答中的代码块以语法高亮形式呈现，支持代码复制和阅读，提升学习体验。"),
        ("对话历史管理", "保留对话上下文，支持连续提问；提供一键清空对话功能，方便重新开始学习。")
    ]
    for name, desc in features:
        p = doc.add_paragraph()
        run = p.add_run(name + "：")
        run.font.bold = True
        p.add_run(desc)

    # ===== 技术原理 =====
    doc.add_heading("技术原理简述", level=1)
    doc.add_paragraph("本项目主要涉及以下核心技术原理：")
    doc.add_paragraph(
        "1. 大语言模型（LLM）API调用原理：大语言模型是基于Transformer架构的深度学习模型，通过海量文本数据训练而成，"
        "能够理解和生成自然语言。本项目通过HTTP请求调用OpenAI兼容的API接口，将用户问题封装为结构化的消息序列"
        "发送给模型，模型返回生成的回答文本。这一过程涉及HTTP POST请求、JSON数据序列化、API认证等计算机网络技术。"
    )
    doc.add_paragraph(
        "2. 本地知识库检索原理：当API不可用时，系统采用关键词匹配的方式在预定义的本地知识库中检索最相关的知识点。"
        "知识库按主题分类存储（如指针、数组、结构体等），每个主题包含关键词列表和对应的知识点内容。"
        "系统对用户输入进行关键词提取和匹配评分，返回得分最高的知识点内容。"
    )
    doc.add_paragraph(
        "3. Web前后端交互原理：项目基于Flask Web框架构建，前端采用HTML+CSS+JavaScript实现用户界面，后端"
        "Python处理业务逻辑。用户通过浏览器发送问题，前端通过Fetch API发起异步HTTP请求，后端处理请求后以JSON格式"
        "返回回答，前端动态更新页面显示。"
    )

    # ===== 实现方案 =====
    doc.add_heading("实现方案与过程", level=1)

    doc.add_heading("1. 使用的工具和平台", level=2)
    tools = [
        ("Python 3", "编程语言，用于后端逻辑开发和知识库构建。"),
        ("Flask 框架", "轻量级Web应用框架，用于构建后端服务和API接口。"),
        ("OpenAI 兼容 API", "大语言模型接口，提供智能对话能力（可配置）。"),
        ("HTML/CSS/JavaScript", "前端技术栈，构建交互式聊天界面。"),
        ("python-docx", "用于生成项目提案报告Word文档。"),
    ]
    for name, desc in tools:
        p = doc.add_paragraph()
        run = p.add_run(name + "：")
        run.font.bold = True
        p.add_run(desc)

    doc.add_heading("2. 关键步骤", level=2)
    steps = [
        "第一步：需求分析与设计。明确项目定位为C语言编程学习助手，设计系统架构为前后端分离的Web应用，确定双模式运行方案。",
        "第二步：知识库构建。梳理C语言核心知识点，涵盖基础语法、控制流、数组、指针、函数、字符串、结构体、文件操作、内存管理、预处理指令、常见错误等11个主题，编写对应的解答内容和代码示例。",
        "第三步：后端开发。使用Flask框架构建Web服务器，实现聊天接口（/chat）、会话重置接口（/reset）和主页路由（/），集成OpenAI API调用和本地知识库检索两个模式。",
        "第四步：前端开发。设计现代风格聊天界面，包含消息展示、代码高亮、打字指示器、对话清空等功能，适配移动端和桌面端。",
        "第五步：测试与优化。验证本地知识库各主题匹配效果，测试API模式下的对话质量，优化前端UI交互体验。",
        "第六步：项目报告撰写。使用python-docx生成完整的项目提案报告，整理项目文档。"
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_heading("3. 运行效果截图", level=2)
    doc.add_paragraph(
        "（注：实际运行时请补充以下截图：\n"
        "1. Web应用首页截图——包含聊天界面和C语言话题标签。\n"
        "2. 对话交互截图——展示用户提问和AI/知识库回答的效果。\n"
        "3. 代码示例展示截图——展示回答中C语言代码高亮的效果。\n"
        "4. 移动端适配截图——展示在手机上的显示效果。）"
    )

    # ===== 总结 =====
    doc.add_heading("项目总结与感悟", level=1)
    doc.add_paragraph(
        "通过本次《人工智能基础与应用》课程项目实践，我深入体验了将AI技术应用于实际问题的完整过程，收获良多。"
    )
    doc.add_paragraph(
        "首先，在技术层面，我学会了如何使用大语言模型的API来构建智能对话应用。从API的调用方式、参数配置、"
        "到消息序列的组织，每一步都需要仔细理解和调试。同时，我也掌握了Flask Web框架的基本使用方法，理解了"
        "前后端交互的完整流程，包括HTTP请求与响应、JSON数据格式、异步通信等计算机网络相关的核心技术。"
    )
    doc.add_paragraph(
        "其次，在项目设计方面，我认识到好的产品设计需要考虑多种场景。本项目设计了双模式运行方案——"
        "在线AI模式和本地知识库模式，既保证了功能的先进性，又确保了基本可用性。这种容错设计在实际工程中具有重要意义。"
    )
    doc.add_paragraph(
        "最后，本次项目也让我认识到AI技术的强大之处和局限。大语言模型能够理解和生成自然语言，为教育领域"
        "带来了革命性的变化。但同时，模型也可能存在知识盲区或生成不准确内容，因此本地知识库作为补充是非常必要的。"
        "未来，我计划继续完善这个项目，增加更多功能，如代码运行环境、练习题生成等，让这个C语言学习小教练"
        "真正成为编程初学者的得力助手。"
    )

    # ===== 保存 =====
    os.makedirs("output", exist_ok=True)
    path = os.path.join("output", "C语言编程学习助手_项目提案报告.docx")
    doc.save(path)
    print("报告已生成: %s" % os.path.abspath(path))
    return path

if __name__ == "__main__":
    generate_report()
